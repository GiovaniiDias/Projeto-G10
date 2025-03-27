from docx import Document

documento = Document('sampleEnchanced.docx')

# como obter o indice + texto
for index, item in enumerate(documento.paragraphs):
    print(index)
    print(item.text)
documento.paragraphs[2].text = 'The population of macacos in alaska is rising'
#documento.save('demo2.docx')

# vai imprimir o index dentro do paragrafo
for index,run in enumerate(documento.paragraphs[4].runs):
    print(str(index) + ' ' + run.text)

# agora faz alteração desejada , conforme o paragafo 4 e os index desejados
documento.paragraphs[4].runs[0].text = 'Um novo paragrafo'
documento.paragraphs[4].runs[1].text = ' texto importante'
documento.paragraphs[4].runs[2].text = 'Outro paragrafo'
documento.paragraphs[6].runs[0].text = 'Onde era o tom got, eu fiz esse '

# alterações para o desafio
documento.paragraphs[10].runs[0].text = 'esse texto foi alterado'
documento.paragraphs[22].text = ' Esse paragrafo oi alterado por completo'
documento.paragraphs[33].runs[0].text = 'Good Morning '

documento.save('demo2.docx')