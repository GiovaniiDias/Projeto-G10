from docx import Document
# add img
from docx.shared import Cm

documento = Document()

# Adicionando um titulo
documento.add_heading('Titulo do Documento', 0)
documento.save('demo.docx')

#add paragrafo
paragrafo = documento.add_paragraph('Um paragrafo simples ')

# add uma nova frase e destacando em 'bold' negrito
paragrafo.add_run('e super importante ').bold = True

# add uma nova frase
paragrafo.add_run('do autor ')

# add uma nova frase e destacando em italico
paragrafo.add_run('Giovani').italic = True

# adicionar hadding(cabeçalho)
documento.add_heading('Titulo nivel 1', level=1)
documento.add_heading('Titulo nivel 2', level=2)
documento.add_heading('Titulo nivel 3', level=3)
documento.add_heading('Titulo nivel 4', level=4)

# formatação de estilo
documento.add_paragraph('Formatação "No Spacing"', style='No Spacing' )
documento.add_paragraph('Formatação "Heading 1"', style='Heading 1' )
documento.add_paragraph('Formatação "Heading 2"', style='Heading 2' )
documento.add_paragraph('Formatação "Heading 3"', style='Heading 3' )
documento.add_paragraph('Formatação "Title"', style='Title' )
documento.add_paragraph('Formatação "Subtitle"', style='Subtitle' )
documento.add_paragraph('Formatação "Quote"', style='Quote' )
documento.add_paragraph('Formatação "Intense Quote"', style='Intense Quote' )
documento.add_paragraph('Formatação "List Paragraph"', style='List Paragraph' )
documento.add_paragraph('Prieiro item em uma lista com pontos', style='List Bullet')
documento.add_paragraph('Primeiro item em uma lista numerada',style='List Number')


# add img
documento.add_picture('nome da img.jpeg', width=Cm(5.25))


documento.save('demo.docx')
