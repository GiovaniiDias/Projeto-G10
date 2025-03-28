from docx import Document

documento = Document()
# criar a lista fora do laço de repetição, para não resetar a lista
mes = input('Mes do Relatório: ')
modelos = []
precos = []
datas = []
vendedores = []
lojas = []

print('Vamos adicionara as vendas agora ')

continuar = 's'
while continuar == 's':
    modelos.append(input('Modelo do produto: '))
    precos.append(input('Preço: '))
    datas.append(input('Data da venda (dd/mm/aaaa): '))
    vendedores.append(input('Vendedor: '))
    lojas.append(input('Loja: '))

    continuar = input('Gostaria de add mais dados? (s/n)')

documento.add_heading('relatorio Mensal de Vendas', 0)
paragrafo = documento.add_paragraph(f'Segue abaixo o relatório de vendas para o mês {mes}')

documento.add_heading(f'Vendas de {mes}',1)

vendas = []
for modelo, preco,data,vendedor,loja in zip(modelos,precos,datas,vendedores,lojas):
    vendas.append([modelo,preco,data,vendedor,loja])


tabela = documento.add_table(rows=1,cols=5)
cabecalho = tabela.rows[0].cells
cabecalho[0].text = 'Modelo'
cabecalho[1].text = 'Preço'
cabecalho[2].text = 'Data'
cabecalho[3].text = 'Vendedor'
cabecalho[4].text = 'Loja'

for modelo,preco,data,vendedor,loja in vendas:
    dados_por_linha = tabela.add_row().cells
    dados_por_linha[0].text = modelo
    dados_por_linha[1].text = preco
    dados_por_linha[2].text = data
    dados_por_linha[3].text = vendedor
    dados_por_linha[4].text = loja

documento.add_paragraph('')

def obter_maior_vendedor(vendas):
    vendedores = []
    for registro in vendas:
        vendedores.append(registro[3])
    vendedores_individuais = set(vendedores)

    maior_vendedor = ''
    ocorrencias = 0

    for vendedor in vendedores_individuais:
        vendas = vendedores.count(vendedor)
        if vendas > ocorrencias:
            maior_vendedor = vendedor
            ocorrencias = vendas
    return maior_vendedor

def obter_loja_com_mais_vendas(vendas):
    lojas = []
    for registro in vendas:
        lojas.append(registro[4])
    lojas_individuais = set(lojas)

    lojas_com_mais_vendas = ''
    ocorrencias = 0

    for loja in lojas_individuais:
        vendas = lojas.count(loja)
        if vendas > ocorrencias:
            lojas_com_mais_vendas = loja
            ocorrencias = vendas
    return lojas_com_mais_vendas


def obter_total_vendas(vendas):
    valores = []
    for registro in vendas:
        valor_normalizado = registro[1].replace (',', '.')
        valores.append(float(valor_normalizado))
    return sum(valores)




maior_vendedor = obter_maior_vendedor(vendas)
loja_com_mais_vendas = obter_loja_com_mais_vendas(vendas)
total_de_vendas = obter_total_vendas(vendas)

documento.add_paragraph(f'Para as vendas deste mês, o funcionario {maior_vendedor} foi o funcionario com maior numero de vendas diretas.')

documento.add_paragraph(f'A loja com mais vendas foi a loja {loja_com_mais_vendas}.')

documento.add_paragraph(f'O lucro total gerado foi de R${total_de_vendas} .')

documento.save('Relatório de vendas.docx')
